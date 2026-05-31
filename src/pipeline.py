from colorsys import yiq_to_rgb
from this import d
import uuid

from src.info_extractor import extract_info
from src.project_classes import Company

from src.doc_generate import generate_project_report,generate_project_notification,generate_project_acceptance_report
import os
import sys

# 获取当前文件的上一级路径
cur_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
class Solver:
    def __init__(self):
        self.sessions = {}
        pass

    def text_preprocess(self, text: str) -> str:
        random_id = uuid.uuid4().hex[:6]
        project_data = extract_info(text)
        project_data.session_id = random_id
        # 将解析结果保存到sessions字典中
        self.sessions[random_id] = project_data
        return random_id
    
    def get_project_data(self, session_id: str) -> dict:
        """
        根据session ID获取保存的项目数据
        
        Args:
            session_id: 会话ID
            
        Returns:
            保存的项目数据字典
        """
        return self.sessions.get(session_id, {"error": "会话ID不存在"})



import json
from src.project_classes import Company, Project

def pipeline(project_data_dict):
    """
    读取project_data.json文件并还原project_data
    
    Returns:
        还原后的Company对象
    """

    # 读取project_data.json文件

    
    # 创建Company对象
    company = Company(
        name=project_data_dict.get("name", ""),
        industry=project_data_dict.get("industry", ""),
        location=project_data_dict.get("location", ""),
        established_year=project_data_dict.get("established_year", 0),
        description=project_data_dict.get("description", ""),
        core_technologies=project_data_dict.get("core_technologies", []),
        patents=project_data_dict.get("patents", []),
        session_id=project_data_dict.get("session_id", "")
    )
    
    # 创建Project对象并添加到Company的projects列表中
    for project_dict in project_data_dict.get("projects", []):
        project = Project(
            name=project_dict.get("name", ""),
            project_id=project_dict.get("project_id", ""),
            company=company,
            start_date=project_dict.get("start_date", ""),
            end_date=project_dict.get("end_date", ""),
            field=project_dict.get("field", ""),
            intellectual_property=project_dict.get("intellectual_property", ""),
            patent_info=project_dict.get("patent_info", ""),
            background=project_dict.get("background", ""),
            objectives=project_dict.get("objectives", []),
            expected_benefits=project_dict.get("expected_benefits", []),
            challenges=project_dict.get("challenges", []),
            budget=project_dict.get("budget", ""),
            team_members=project_dict.get("team_members", [])
        )
        company.projects.append(project)
        
    

    ## 生成立项报告
    # project_reports = []
    # for project in company.projects:
    #     doc_path = process_one_project(project)
    #     project_reports.append(doc_path)
    return _pipeline(company)


def _pipeline(company):
    import concurrent.futures
    import os
    from src.retry_utils import retry_with_backoff
    
    # 指定进程数量（例如：使用CPU核心数的一半）
    cpu_count = os.cpu_count()
    # max_workers = max(1, cpu_count // 2)
    max_workers = 3
    print(f'cpu_count : {cpu_count}')
    
    # 创建带重试机制的项目处理函数
    process_with_retry = retry_with_backoff(max_retries=5)(process_one_project)
    
    # 使用多进程并行处理项目
    project_reports = []
    with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
        project_reports = list(executor.map(process_with_retry, company.projects))

    # for project in company.projects:
    #     doc_path = process_with_retry(project)
    #     project_reports.append(doc_path)

    return project_reports


def process_one_project(project: Project) -> dict:
    """
    处理一个项目，生成项目报告
    
    Args:
        project: 项目对象
        
    Returns:
        项目报告结构化数据
    """
    project_report = generate_project_report(project)
    project.report = project_report
    project_notification = generate_project_notification(project)
    project.notification = project_notification
    project_acceptance_report = generate_project_acceptance_report(project)
    project.acceptance_report = project_acceptance_report


    gen_final_docx(project)


def print_runs(runs):

    for idx, run in enumerate(runs):
        print(f'run {idx}: {run.text}')


def gen_final_docx(project):
        # 保存生成的三个报告用于测试
    # with open(f'./assets/test_report.json', 'w', encoding='utf-8') as f:
    #     json.dump(project.report, f, ensure_ascii=False, indent=4)
    # with open(f'./assets/test_notification.json', 'w', encoding='utf-8') as f:
    #     json.dump(project.notification, f, ensure_ascii=False, indent=4)
    # with open(f'./assets/test_acceptance_report.json', 'w', encoding='utf-8') as f:
    #     json.dump(project.acceptance_report, f, ensure_ascii=False, indent=4)

    project_report = project.report
    project_notification = project.notification
    project_acceptance_report = project.acceptance_report


    doc_template_path = './assets/templates/1.docx'

    from docx import Document
    doc = Document(doc_template_path)

    texts = [ para.text for para in doc.paragraphs]
    # for para in doc.paragraphs:
    #     if '项目名称' in para.text:
    #         para.text = para.text.replace('项目名称', project.name)
    ## 逐行替换
    # 替换项目名称
    # doc.paragraphs[4].text = doc.paragraphs[4].text.replace('高精准金型模具的研发', project.name)

    doc.paragraphs[4].runs[2].text = project.name
    # 替换企业名称
    # doc.paragraphs[5].text = doc.paragraphs[5].text.replace('上海索屿智能科技有限公司', project.company.name)
    doc.paragraphs[5].runs[2].text = project.company.name

    doc.paragraphs[6].runs[2].text = ''

    # 替换企业法人
    # doc.paragraphs[6].text = doc.paragraphs[6].text.replace('郑  涵', '')
    # 替换项目负责人
    # doc.paragraphs[7].text = doc.paragraphs[7].text.replace('汪永太', '')
    doc.paragraphs[7].runs[2].text = ''
    # 替换项目起止时间
    start_year = project.start_date.split('.')[0]
    start_month = project.start_date.split('.')[1].zfill(2)

    doc.paragraphs[8].runs[1].text = doc.paragraphs[8].runs[1].text.replace('2024', start_year)
    doc.paragraphs[8].runs[3].text = doc.paragraphs[8].runs[3].text.replace('01', start_month)
    end_year = project.end_date.split('.')[0]
    end_month = project.end_date.split('.')[1].zfill(2)
    doc.paragraphs[8].runs[5].text = doc.paragraphs[8].runs[5].text.replace('2024', end_year)
    doc.paragraphs[8].runs[7].text = doc.paragraphs[8].runs[7].text.replace('12', end_month)

    doc.paragraphs[14].runs[1].text = doc.paragraphs[14].runs[1].text.replace('2024', start_year)
    doc.paragraphs[14].runs[3].text = doc.paragraphs[14].runs[3].text.replace('01', start_month)

    # 项目背景
  
    doc.paragraphs[19].runs[0].text = project_report['项目背景'][0].strip('。.')

    doc.paragraphs[20].runs[0].text = project_report['项目背景'][1]
    if not doc.paragraphs[20].runs[0].text.endswith('。'):
        doc.paragraphs[20].runs[0].text += '。'

    doc.paragraphs[21].runs[0].text = project_report['项目背景'][2].strip('。.')

    doc.paragraphs[23].text = project_report['项目目标']['目标简介']

    for target_idx in range(4):
        doc.paragraphs[24 + target_idx].runs[0].text = project_report['项目目标']['目标列表'][target_idx]


    # 预期效益
    doc.paragraphs[30].text = project_report['预期效益与挑战']['预期效益']['经济效益']
    doc.paragraphs[31].text = project_report['预期效益与挑战']['预期效益']['社会效益']
    doc.paragraphs[32].text = project_report['预期效益与挑战']['预期效益']['技术效益']


    # 面临的挑战
    doc.paragraphs[34].text = project_report['预期效益与挑战']['面临的挑战'][0]
    doc.paragraphs[35].text = project_report['预期效益与挑战']['面临的挑战'][1]
    doc.paragraphs[36].text = project_report['预期效益与挑战']['面临的挑战'][2]
    doc.paragraphs[37].text = project_report['预期效益与挑战']['面临的挑战'][3]


    doc.paragraphs[39].runs[0].text = project.company.name

    
    doc.paragraphs[40].runs[0].text = start_year


    doc.paragraphs[40].runs[2].text = start_month[0]
    doc.paragraphs[40].runs[3].text = start_month[1]



    doc.paragraphs[43].runs[0].text = project_notification['通知正文']
    # 立项通知
    doc.paragraphs[46].runs[1].text = ''

    doc.paragraphs[47].runs[1].text = start_year
    doc.paragraphs[47].runs[3].text = start_month[0]
    doc.paragraphs[47].runs[4].text = start_month[1]

    doc.paragraphs[48].runs[0].text = '项目名称：' + project.name
    
    doc.paragraphs[49].runs[1].text = ""

    doc.paragraphs[50].runs[1].text = ""

    doc.paragraphs[51].runs[1].text = start_year
    doc.paragraphs[51].runs[3].text = start_month[0]
    doc.paragraphs[51].runs[4].text = start_month[1]

    doc.paragraphs[51].runs[6].text = end_year
    doc.paragraphs[51].runs[8].text = end_month[1]



    doc.paragraphs[59].runs[0].text = project.company.name


    doc.paragraphs[60].runs[0].text = start_year
    doc.paragraphs[60].runs[2].text = start_month[0]
    doc.paragraphs[60].runs[3].text = start_month[1]



    # 替换验收报告部分
    doc.paragraphs[63].text = project_acceptance_report['项目概况']
    doc.paragraphs[65].text = project_acceptance_report['项目目标与完成情况']
    for i, tech_point in enumerate(project_acceptance_report['项目核心技术点与创新点']['核心技术点'], 68):
        doc.paragraphs[i].text = tech_point
    for i, innovation in enumerate(project_acceptance_report['项目核心技术点与创新点']['创新点'], 75):
        doc.paragraphs[i].text = innovation

    doc.paragraphs[82].text = project_acceptance_report['项目验收情况']['验收内容']['验收简介']

    doc.paragraphs[85].text = "验收方式：" +project_acceptance_report['项目验收情况']['验收内容']['知识产权验收']['验收方式']

    doc.paragraphs[86].text = "验收成果：" +project_acceptance_report['项目验收情况']['验收内容']['知识产权验收']['验收成果']

    # 修改验收内容

    start_idx = 3

    while start_idx < 15:
        x = start_idx // 3 -1
        y = start_idx % 3

        key = None
        if y == 0:
            key = f'核心指标项{x+1}'
        elif y == 1:
            key = '目标值/验收值'
        elif y == 2:
            key = '实测效果描述'

        doc.tables[1]._cells[start_idx].text = project_acceptance_report['项目验收情况']['验收内容']['核心技术指标达成情况验收'][x][key]
        
        start_idx += 1



    doc.paragraphs[88].text = project_acceptance_report['项目验收情况']['验收结论']
    

    doc.paragraphs[90].text = project.company.name

    doc.paragraphs[91].runs[0].text = end_year
    doc.paragraphs[91].runs[2].text = end_month


    # 修改眉页
    # 访问文档的页眉
    header = doc.sections[0].header
    # 遍历页眉中的段落，替换公司名称
    for paragraph in header.paragraphs:
        if '上海骧和信息技术有限公司' in paragraph.text:
            paragraph.text = paragraph.text.replace('上海骧和信息技术有限公司', project.company.name)



    # 保存修改后的文档
    import datetime
    timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
    # 保存到上一级的assets/docs目录下
    save_dir = os.path.join(cur_dir, 'assets', 'docs')
    # 确保目录存在
    os.makedirs(save_dir, exist_ok=True)
    save_path = os.path.join(save_dir, f'{project.name}_{timestamp}.docx')
    doc.save(save_path)

    return project_report





    

    

    